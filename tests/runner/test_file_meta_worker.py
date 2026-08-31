"""Tests for the file metadata worker (plan 039 Phase 0, task 0.3).

Real Pillow + pypdf/python-docx are exercised (the fileintel extra is
installed in the dev venv); EXIF-bearing images are generated with
Pillow itself, so no binary fixtures are committed.
"""

from __future__ import annotations

import os
from pathlib import Path

import pytest

# The worker's extractors need the optional fileintel extra; skip the
# whole module cleanly when it's absent (same posture as the dictation
# engine's sherpa smoke test).
pytest.importorskip("PIL", reason="fileintel extra (pillow) not installed")

from agent_meow.entities.file_index import (
    KIND_DOCUMENT,
    KIND_IMAGE,
    KIND_OTHER,
    STATUS_GONE,
    STATUS_INDEXED,
    classify_kind,
)
from agent_meow.runner.file_meta_worker import (
    compute_dhash,
    compute_doc_hash,
    extract_document_meta,
    extract_image_meta,
    process_entry,
    run_once,
)
from agent_meow.stores.file_index_store.sqlalchemy_store import (
    SqlAlchemyFileIndexStore,
)


def _make_store(tmp_path: Path) -> SqlAlchemyFileIndexStore:
    return SqlAlchemyFileIndexStore(f"sqlite:///{tmp_path / 'idx.db'}")


def _enqueue(store, tmp_path: Path, name: str, data: bytes) -> str:
    ws = tmp_path / "ws"
    ws.mkdir(exist_ok=True)
    f = ws / name
    f.write_bytes(data)
    return store.upsert_pending(
        host_id="h1",
        workspace=str(ws),
        path=str(f),
        kind=classify_kind(str(f)),
        size=len(data),
        mtime_ns=f.stat().st_mtime_ns,
    )


def _png_bytes(w: int = 32, h: int = 32, color: tuple[int, int, int] = (255, 0, 0)) -> bytes:
    import io

    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (w, h), color).save(buf, "PNG")
    return buf.getvalue()


def _gradient_png(seed: int, size: int = 64) -> bytes:
    """A structured image dHash can distinguish (solid colors all hash to 0)."""
    import io
    import random

    from PIL import Image

    rnd = random.Random(seed)
    img = Image.new("L", (size, size))
    px = img.load()
    for y in range(size):
        for x in range(size):
            px[x, y] = (x * 4 + y * 2 + seed * 7 + rnd.randint(0, 3)) % 256
    buf = io.BytesIO()
    img.save(buf, "PNG")
    return buf.getvalue()


def _jpeg_with_exif(tmp_path: Path) -> Path:
    """A JPEG carrying EXIF (camera, DateTimeOriginal, GPS) built with Pillow."""
    from PIL import Image

    img = Image.new("RGB", (64, 48), (10, 20, 30))
    exif = img.getexif()
    exif[0x010F] = "Canon"  # Make
    exif[0x0110] = "EOS R5"  # Model
    exif[0x0132] = "2026:08:30 12:00:00"  # DateTime
    gps_ifd = exif.get_ifd(0x8825)
    # 22°18'0"N 114°12'0"E (Shenzhen-ish)
    gps_ifd[1] = "N"
    gps_ifd[2] = (22.0, 18.0, 0.0)
    gps_ifd[3] = "E"
    gps_ifd[4] = (114.0, 12.0, 0.0)
    out = tmp_path / "exif.jpg"
    img.save(out, "JPEG", exif=exif)
    return out


# ── pure extractors ──────────────────────────────────────────────────────────


def test_dhash_stable_across_reencode_and_differs_for_different_images(tmp_path: Path):
    from PIL import Image

    a = tmp_path / "a.png"
    a.write_bytes(_gradient_png(1))
    # Resize + re-encode (same content) → identical dHash.
    b = tmp_path / "b.jpg"
    Image.open(a).resize((32, 32)).save(b, "JPEG", quality=70)
    assert compute_dhash(str(a)) == compute_dhash(str(b))

    c = tmp_path / "c.png"
    c.write_bytes(_gradient_png(99))
    assert compute_dhash(str(a)) != compute_dhash(str(c))


def test_extract_image_meta_reads_exif_fields(tmp_path: Path):
    path = _jpeg_with_exif(tmp_path)
    meta = extract_image_meta(str(path))
    assert meta["camera_make"] == "Canon"
    assert meta["camera_model"] == "EOS R5"
    assert meta["exif_datetime"] == "2026:08:30 12:00:00"
    assert meta["width"] == 64 and meta["height"] == 48
    # GPS converted DMS → decimal.
    assert abs(meta["gps_lat"] - 22.3) < 1e-4
    assert abs(meta["gps_lon"] - 114.2) < 1e-4
    assert "dhash" in meta and len(meta["dhash"]) == 16
    assert meta["thumb_path"] and os.path.isfile(meta["thumb_path"])


def test_extract_image_meta_handles_no_exif(tmp_path: Path):
    f = tmp_path / "plain.png"
    f.write_bytes(_png_bytes())
    meta = extract_image_meta(str(f))
    assert meta["width"] == 32
    assert "camera_make" not in meta
    assert "dhash" in meta


def test_extract_document_meta_pdf_and_docx_and_txt(tmp_path: Path):
    # txt
    t = tmp_path / "n.txt"
    t.write_text("发票 invoice number 42 你好", encoding="utf-8")
    m = extract_document_meta(str(t))
    assert m["words"] >= 4
    assert "发票" in m["text_excerpt"]

    # docx
    d = tmp_path / "doc.docx"
    import docx

    doc = docx.Document()
    doc.add_paragraph("quarterly report summary")
    doc.add_paragraph("second line")
    doc.save(d)
    m2 = extract_document_meta(str(d))
    assert m2["paragraphs"] == 2
    assert "quarterly report" in m2["text_excerpt"]


def test_doc_hash_detects_identical_content(tmp_path: Path):
    a = tmp_path / "a.pdf"
    b = tmp_path / "b.pdf"
    payload = b"%PDF-1.4 identical"
    a.write_bytes(payload)
    b.write_bytes(payload)
    assert compute_doc_hash(str(a)) == compute_doc_hash(str(b))
    c = tmp_path / "c.pdf"
    c.write_bytes(b"%PDF-1.4 other")
    assert compute_doc_hash(str(a)) != compute_doc_hash(str(c))


# ── process_entry / run_once ────────────────────────────────────────────────


def test_process_entry_indexes_image_with_meta(tmp_path: Path):
    store = _make_store(tmp_path)
    fid = _enqueue(store, tmp_path, "photo.jpg", _jpeg_with_exif(tmp_path).read_bytes())
    entry = store.claim_pending()[0]
    assert entry.id == fid
    outcome = process_entry(store, entry)
    assert outcome == "indexed"
    row = store.find_by_path(
        host_id="h1", workspace=str(tmp_path / "ws"), path=str(tmp_path / "ws" / "photo.jpg")
    )
    assert row.status == STATUS_INDEXED
    assert row.content_hash and len(row.content_hash) == 16  # dhash
    assert row.meta["camera_make"] == "Canon"
    assert row.thumb_path and os.path.isfile(row.thumb_path)


def test_duplicate_image_content_marked_duplicate(tmp_path: Path):
    store = _make_store(tmp_path)
    data = _jpeg_with_exif(tmp_path).read_bytes()
    _enqueue(store, tmp_path, "orig.jpg", data)
    assert run_once(store)["indexed"] == 1
    _enqueue(store, tmp_path, "copy.jpg", data)
    assert run_once(store)["duplicate"] == 1
    row = store.find_by_path(
        host_id="h1", workspace=str(tmp_path / "ws"), path=str(tmp_path / "ws" / "copy.jpg")
    )
    assert row.status == "duplicate"
    assert row.content_hash  # hash recorded for provenance


def test_process_entry_parks_undecodable_image(tmp_path: Path):
    store = _make_store(tmp_path)
    # .jpg extension but not a real image → dhash None → byte hash, still indexes.
    _enqueue(store, tmp_path, "broken.jpg", b"not an image at all")
    entry = store.claim_pending()[0]
    outcome = process_entry(store, entry)
    assert outcome == "indexed"
    row = store.find_by_path(
        host_id="h1", workspace=str(tmp_path / "ws"), path=str(tmp_path / "ws" / "broken.jpg")
    )
    assert len(row.content_hash) == 32  # sha256 fallback


def test_process_entry_marks_gone_for_missing_file(tmp_path: Path):
    store = _make_store(tmp_path)
    _enqueue(store, tmp_path, "ghost.jpg", b"x")
    os.remove(tmp_path / "ws" / "ghost.jpg")
    entry = store.claim_pending()[0]
    assert process_entry(store, entry) == "gone"
    row = store.find_by_path(
        host_id="h1", workspace=str(tmp_path / "ws"), path=str(tmp_path / "ws" / "ghost.jpg")
    )
    assert row.status == STATUS_GONE


def test_run_once_processes_mixed_queue(tmp_path: Path):
    store = _make_store(tmp_path)
    _enqueue(store, tmp_path, "a.jpg", _jpeg_with_exif(tmp_path).read_bytes())
    _enqueue(store, tmp_path, "b.txt", b"hello world")
    _enqueue(store, tmp_path, "c.bin", b"\x00\x01")
    tally = run_once(store, batch=8)
    assert tally.get("indexed", 0) == 3
    rows = store.list_workspace(host_id="h1", workspace=str(tmp_path / "ws"))
    kinds = {Path(r.path).name: r.kind for r in rows}
    assert kinds["a.jpg"] == KIND_IMAGE
    assert kinds["b.txt"] == KIND_DOCUMENT
    assert kinds["c.bin"] == KIND_OTHER


def test_worker_survives_extractor_crash(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    """A raising extractor parks the row failed; the batch continues."""
    store = _make_store(tmp_path)
    _enqueue(store, tmp_path, "boom.jpg", _png_bytes())
    _enqueue(store, tmp_path, "ok.txt", b"fine")
    import agent_meow.runner.file_meta_worker as w

    real = w.extract_image_meta

    def _boom(path: str) -> dict:
        if "boom" in path:
            raise RuntimeError("simulated crash")
        return real(path)

    monkeypatch.setattr(w, "extract_image_meta", _boom)
    tally = run_once(store, batch=8)
    assert tally.get("failed") == 1
    assert tally.get("indexed") == 1
    ws = str(tmp_path / "ws")
    failed = store.find_by_path(host_id="h1", workspace=ws, path=str(tmp_path / "ws" / "boom.jpg"))
    assert failed.status == "failed"
    assert "simulated crash" in (failed.error or "")
