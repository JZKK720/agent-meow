"""File metadata worker — deterministic EXIF / text / hash extraction (plan 039).

Consumes the ``file_index`` queue the watcher fills: claims
``pending`` rows, extracts metadata **in-process with no LLM** (the
2026-08-29 lesson: the old pipeline posted chat messages to the vision
agent and hoped it cooperated), then marks each row ``indexed``,
``duplicate`` (content hash already owned), or parks it ``failed``.

Extracted per kind:

- image: Pillow EXIF (capture date, GPS lat/lon, camera make/model,
  lens, dimensions, orientation), a 64-bit **dHash** (content hash for
  dedup — PhotoPrism's perceptual-hash pattern, hand-rolled to avoid an
  imagehash dep), and a WebP thumbnail under the data dir.
- document: text excerpt + page/word counts via pypdf / python-docx /
  plain read; sha256-of-first-1MB as the content hash.
- other: size/mtime only, sha256 content hash.

Everything imports the optional deps lazily so the module loads without
the ``fileintel`` extra (the runner logs and no-ops instead).
"""

from __future__ import annotations

import hashlib
import logging
import os
import threading
from collections.abc import Sequence
from pathlib import Path
from typing import TYPE_CHECKING, cast

from agent_meow.entities.file_index import (
    KIND_DOCUMENT,
    KIND_IMAGE,
    FileIndexEntry,
)
from agent_meow.stores.file_index_store import FileIndexStore

if TYPE_CHECKING:
    # Pillow ships inline types (py.typed); Image.Exif is the real
    # boundary type for EXIF access.
    from PIL import Image

_logger = logging.getLogger(__name__)

# Text excerpt length stored for FTS (plan 039 phase 1) and previews.
_TEXT_EXCERPT_CHARS = 4000
# Thumbnail longest edge (px).
_THUMB_EDGE = 256
# Claim batch size per cycle.
_DEFAULT_BATCH = 8


def _thumbs_dir() -> Path:
    """Thumbnail store: <data_dir>/file_thumbs (created on demand)."""
    from agent_meow.host.local_server import _local_data_dir

    d = _local_data_dir() / "file_thumbs"
    d.mkdir(parents=True, exist_ok=True)
    return d


# ── image metadata ───────────────────────────────────────────────────────────


def _exif_to_dict(exif: Image.Exif) -> dict[str, object]:
    """Pull the fields we index out of a Pillow Exif object."""
    out: dict[str, object] = {}
    try:
        if d := exif.get(0x0100):  # ImageWidth
            out["width"] = int(d)
        if d := exif.get(0x0101):  # ImageLength
            out["height"] = int(d)
        if d := exif.get(0x011A):  # XResolution
            out["x_resolution"] = float(d)
        if d := exif.get(0x0132):  # DateTime
            out["exif_datetime"] = str(d)
        if d := exif.get(0x010F):  # Make
            out["camera_make"] = str(d).strip()
        if d := exif.get(0x0110):  # Model
            out["camera_model"] = str(d).strip()
        if d := exif.get(0x0112):  # Orientation
            out["orientation"] = int(d)
        # Sub-IFDs: EXIF (0x8769) and GPS (0x8825).
        try:
            sub = exif.get_ifd(0x8769)
            if d := sub.get(0x9003):  # DateTimeOriginal
                out["datetime_original"] = str(d)
            if d := sub.get(0xA434):  # LensModel
                out["lens_model"] = str(d).strip()
        except Exception:  # noqa: BLE001
            pass
        try:
            gps = exif.get_ifd(0x8825)
            # get_ifd hands back a dict[int, Any] from Pillow's untyped
            # internals; the values are IFDRational/str we only read via
            # str()/float() conversions, so widening the key/value types
            # here is safe.
            coords = _gps_decimal(cast("dict[object, object]", gps))
            if coords is not None:
                out["gps_lat"], out["gps_lon"] = coords
        except Exception:  # noqa: BLE001
            pass
    except Exception:  # noqa: BLE001 — a malformed EXIF block must not kill indexing
        _logger.debug("exif parse failed", exc_info=True)
    return out


def _dms_to_decimal(value: object, ref: str) -> float | None:
    """Convert EXIF degrees/minutes/seconds rationals + N/S/E/W ref to decimal."""
    try:
        if not isinstance(value, Sequence) or isinstance(value, (str, bytes)):
            return None
        deg, minutes, sec = (float(v) for v in value)
    except (TypeError, ValueError):
        return None
    dec = deg + minutes / 60.0 + sec / 3600.0
    if str(ref).upper() in {"S", "W"}:
        dec = -dec
    return round(dec, 6)


def _gps_decimal(gps: dict[object, object]) -> tuple[float, float] | None:
    lat = _dms_to_decimal(gps.get(2), str(gps.get(1, "N")))
    lon = _dms_to_decimal(gps.get(4), str(gps.get(3, "E")))
    if lat is None or lon is None:
        return None
    return lat, lon


def compute_dhash(path: str) -> str | None:
    """64-bit difference hash (16 hex chars) of a grayscale 8x9 resize.

    The content hash for images: a re-saved copy (same pixels) matches;
    a resize/crop/compress-tweak stays close in Hamming distance.
    """
    try:
        from PIL import Image
    except ImportError:  # pragma: no cover — fileintel extra absent
        return None
    try:
        with Image.open(path) as img:
            g = img.convert("L").resize((9, 8))
            px = list(g.getdata())
            bits = 0
            for row in range(8):
                for col in range(8):
                    left = px[row * 9 + col]
                    right = px[row * 9 + col + 1]
                    bits = (bits << 1) | (1 if left > right else 0)
            return f"{bits:016x}"
    except Exception:  # noqa: BLE001 — undecodable image
        return None


def extract_image_meta(path: str) -> dict[str, object]:
    """EXIF + dimensions + dHash + thumbnail for one image file.

    Returns a dict with keys: kind fields plus ``dhash`` and
    ``thumb_path`` (the worker persists those into the index row).
    An undecodable image (corrupt download) yields a sparse dict — the
    row still indexes by name/size so the panel + FTS see it, matching
    the document path.
    """
    meta: dict[str, object] = {}
    try:
        from PIL import Image, ImageOps
    except ImportError:
        raise RuntimeError("pillow not installed — pip install agent-meow[fileintel]") from None
    try:
        with Image.open(path) as img:
            meta["format"] = img.format or Path(path).suffix.lstrip(".").upper()
            meta["width"], meta["height"] = img.size
            exif = img.getexif()
            meta.update(_exif_to_dict(exif))
            # Thumbnail: auto-rotated by EXIF orientation, longest edge 256.
            try:
                thumb = ImageOps.exif_transpose(img)
                thumb.thumbnail((_THUMB_EDGE, _THUMB_EDGE))
                digest = hashlib.sha256(path.encode()).hexdigest()[:16]
                thumb_path = _thumbs_dir() / f"{digest}.webp"
                if not thumb_path.exists():
                    thumb.convert("RGB").save(thumb_path, "WEBP", quality=80)
                meta["thumb_path"] = str(thumb_path)
            except Exception:  # noqa: BLE001 — a bad thumbnail must not fail indexing
                _logger.debug("thumbnail failed for %s", path, exc_info=True)
    except Exception:  # noqa: BLE001 — undecodable image: index by name, no pixels
        _logger.debug("image decode failed for %s", path, exc_info=True)
        meta.pop("width", None)
        meta.pop("height", None)
    if (dh := compute_dhash(path)) is not None:
        meta["dhash"] = dh
    return meta


# ── document metadata ────────────────────────────────────────────────────────


def extract_document_meta(path: str) -> dict[str, object]:
    """Page/word counts + a text excerpt for PDF/DOCX/plain documents."""
    meta: dict[str, object] = {}
    ext = Path(path).suffix.lower()
    text: str | None = None
    try:
        if ext == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(path)
            meta["pages"] = len(reader.pages)
            text = " ".join((page.extract_text() or "") for page in reader.pages[:10])
        elif ext == ".docx":
            import docx

            d = docx.Document(path)
            text = "\n".join(p.text for p in d.paragraphs)
            meta["paragraphs"] = len(d.paragraphs)
        else:
            # txt/md/csv — plain read, capped.
            with open(path, encoding="utf-8", errors="replace") as fh:
                text = fh.read(_TEXT_EXCERPT_CHARS * 2)
    except Exception:  # noqa: BLE001 — unreadable doc still indexes by name
        _logger.debug("doc text extract failed for %s", path, exc_info=True)
    if text:
        words = text.split()
        meta["words"] = len(words)
        meta["text_excerpt"] = text[:_TEXT_EXCERPT_CHARS]
    return meta


def compute_doc_hash(path: str) -> str:
    """sha256 of the first 1 MB — dedup key for non-image files."""
    h = hashlib.sha256()
    with open(path, "rb") as fh:
        h.update(fh.read(1024 * 1024))
    return h.hexdigest()[:32]


# ── the worker ───────────────────────────────────────────────────────────────


def process_entry(store: FileIndexStore, entry: FileIndexEntry) -> str:
    """Extract + persist metadata for one claimed row. Returns outcome.

    Outcomes: ``indexed`` | ``duplicate`` | ``failed``.
    """
    path = entry.path
    try:
        if not os.path.isfile(path):
            store.mark_gone(host_id=entry.host_id, workspace=entry.workspace, path=path)
            return "gone"
        if entry.kind == KIND_IMAGE:
            meta = extract_image_meta(path)
            content_hash = str(meta.pop("dhash", "") or "")
            raw_thumb = meta.pop("thumb_path", None)
            thumb_path = str(raw_thumb) if raw_thumb else None
            if not content_hash:
                content_hash = compute_doc_hash(path)  # non-decodable image → byte hash
        elif entry.kind == KIND_DOCUMENT:
            meta = extract_document_meta(path)
            content_hash = compute_doc_hash(path)
            thumb_path = None
        else:
            meta = {}
            content_hash = compute_doc_hash(path)
            thumb_path = None

        if content_hash:
            owner = store.find_hash_owner(
                host_id=entry.host_id,
                workspace=entry.workspace,
                content_hash=content_hash,
                exclude_id=entry.id,
            )
            if owner is not None:
                store.mark_duplicate(entry.id, content_hash=content_hash)
                return "duplicate"

        store.mark_indexed(
            entry.id,
            content_hash=content_hash,
            meta=meta,
            thumb_path=thumb_path,
        )
        return "indexed"
    except Exception as exc:  # noqa: BLE001 — one bad file must not kill the batch
        _logger.warning("file index worker failed for %s: %s", path, exc)
        store.mark_failed(entry.id, f"{type(exc).__name__}: {exc}")
        return "failed"


def run_once(store: FileIndexStore, *, batch: int = _DEFAULT_BATCH) -> dict[str, int]:
    """Claim and process one batch. Returns an outcome histogram.

    The runner's asyncio loop calls this on a worker thread (Pillow/pypdf
    are blocking), so it never touches the event loop.
    """
    claimed = store.claim_pending(limit=batch)
    tally: dict[str, int] = {}
    for entry in claimed:
        outcome = process_entry(store, entry)
        tally[outcome] = tally.get(outcome, 0) + 1
    return tally


def worker_loop(
    store: FileIndexStore,
    *,
    stop_event: threading.Event,
    interval: float = 2.0,
    batch: int = _DEFAULT_BATCH,
) -> None:
    """Poll-and-process loop until ``stop_event`` is set.

    The queue itself is durable (rows stay ``pending`` in SQLite), so
    this is a *consumer* poll, not the old filesystem poll: latency is
    bounded by ``interval`` only when events are missed, and watchdog
    events + the bootstrap scan do the real enqueueing.
    """
    while not stop_event.is_set():
        try:
            tally = run_once(store, batch=batch)
            if (
                any(tally.values())
                and (tally.get("indexed", 0) + tally.get("duplicate", 0)) >= batch
            ):
                continue  # backlog remains — drain immediately
        except Exception:  # noqa: BLE001 — store hiccups must not kill the thread
            _logger.warning("file index worker cycle failed", exc_info=True)
        stop_event.wait(interval)
